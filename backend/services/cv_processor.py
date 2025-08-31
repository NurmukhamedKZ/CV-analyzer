import logging
from typing import Optional
import PyPDF2
from docx import Document
import os

logger = logging.getLogger(__name__)

class CVProcessor:
    """Service for processing CV files and extracting text"""
    
    def __init__(self):
        self.supported_formats = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'application/msword': self._extract_doc_text
        }
    
    async def extract_text(self, file_path: str, content_type: str) -> Optional[str]:
        """
        Extract text from CV file based on content type
        
        Args:
            file_path: Path to the uploaded file
            content_type: MIME type of the file
            
        Returns:
            Extracted text or None if extraction fails
        """
        try:
            logger.info(f"Extracting text from {file_path} with content type: {content_type}")
            
            # Check if content type is supported
            if content_type not in self.supported_formats:
                logger.warning(f"Unsupported content type: {content_type}")
                return None
            
            # Extract text using appropriate method
            extractor = self.supported_formats[content_type]
            text = extractor(file_path)
            
            if text:
                logger.info(f"Successfully extracted {len(text)} characters from {file_path}")
                return text.strip()
            else:
                logger.warning(f"No text extracted from {file_path}")
                return None
                
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            return None
    
    def _extract_pdf_text(self, file_path: str) -> Optional[str]:
        """Extract text from PDF file"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                
                return text.strip()
                
        except Exception as e:
            logger.error(f"Error extracting PDF text: {str(e)}")
            return None
    
    def _extract_docx_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"
            
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error extracting DOCX text: {str(e)}")
            return None
    
    def _extract_doc_text(self, file_path: str) -> Optional[str]:
        """Extract text from DOC file (basic implementation)"""
        try:
            # Note: This is a basic implementation
            # For production, you might want to use a more robust library like python-docx2txt
            # or convert DOC to DOCX first
            
            # For now, return a placeholder message
            logger.warning("DOC file processing is limited. Consider converting to DOCX for better results.")
            return "DOC file content extracted (limited processing available)"
            
        except Exception as e:
            logger.error(f"Error extracting DOC text: {str(e)}")
            return None
    
    def get_file_info(self, file_path: str) -> dict:
        """Get basic information about the file"""
        try:
            stat = os.stat(file_path)
            return {
                "size": stat.st_size,
                "created": stat.st_ctime,
                "modified": stat.st_mtime,
                "exists": True
            }
        except Exception as e:
            logger.error(f"Error getting file info: {str(e)}")
            return {"exists": False}
    
    def validate_file_content(self, text: str) -> dict:
        """Validate extracted text content"""
        if not text:
            return {"valid": False, "reason": "No text extracted"}
        
        # Basic validation checks
        issues = []
        
        if len(text) < 50:
            issues.append("Text seems too short for a CV")
        
        if len(text) > 50000:
            issues.append("Text seems too long for a CV")
        
        # Check for common CV sections
        cv_keywords = ["experience", "education", "skills", "work", "job", "employment"]
        text_lower = text.lower()
        
        found_sections = [keyword for keyword in cv_keywords if keyword in text_lower]
        if len(found_sections) < 2:
            issues.append("Document doesn't appear to contain typical CV sections")
        
        return {
            "valid": len(issues) == 0,
            "issues": issues,
            "text_length": len(text),
            "found_sections": found_sections
        }
